"""
Resume Generator for Charles Martin, Ed.D.
Matches the format of the Natalie Anz reference resume
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# Color constants (matching reference resume)
BLUE_ACCENT = RGBColor(91, 155, 213)  # #5B9BD5 - name, company names, degree names
BLUE_LIGHT_BG = "DEEAF6"  # Section header background
BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(102, 102, 102)

def set_cell_shading(cell, color_hex):
    """Set cell background color"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_horizontal_line(doc):
    """Add a thin horizontal line"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # Use a bottom border on the paragraph
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '<w:bottom w:val="single" w:sz="6" w:space="1" w:color="000000"/>'
        '</w:pBdr>'
    )
    p._p.get_or_add_pPr().append(pBdr)

def add_section_header(doc, text):
    """Add a section header with blue background bar"""
    # Create a single-cell table for the background
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, BLUE_LIGHT_BG)

    # Add the header text
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # Set table width to full page
    table.autofit = False
    for cell in table.rows[0].cells:
        cell.width = Inches(7.0)

    # Add spacing
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def add_job_entry(doc, title, company, dates, bullets):
    """Add a job entry with title, company, dates, and bullets"""
    # Title and dates on same line using a table
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False

    # Title cell (left)
    title_cell = table.cell(0, 0)
    title_cell.width = Inches(5.0)
    p = title_cell.paragraphs[0]
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # Dates cell (right-aligned)
    dates_cell = table.cell(0, 1)
    dates_cell.width = Inches(2.0)
    p = dates_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(dates)
    run.italic = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # Remove table borders
    for row in table.rows:
        for cell in row.cells:
            set_cell_borders(cell, None)

    # Company name in blue
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(company)
    run.font.color.rgb = BLUE_ACCENT
    run.font.size = Pt(10)
    run.font.name = 'Calibri'

    # Bullet points
    for bullet in bullets:
        p = doc.add_paragraph(bullet, style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = 'Calibri'

    # Add spacing after job entry
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def set_cell_borders(cell, val):
    """Remove cell borders"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '<w:top w:val="nil"/>'
        '<w:left w:val="nil"/>'
        '<w:bottom w:val="nil"/>'
        '<w:right w:val="nil"/>'
        '</w:tcBorders>'
    )
    tcPr.append(tcBorders)

def add_education_entry(doc, degree, institution, in_progress=False):
    """Add an education entry with degree in blue"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)

    # Degree name in bold blue
    run = p.add_run(degree)
    run.bold = True
    run.font.color.rgb = BLUE_ACCENT
    run.font.size = Pt(10)
    run.font.name = 'Calibri'

    if in_progress:
        run2 = p.add_run(" (in progress)")
        run2.font.color.rgb = BLUE_ACCENT
        run2.font.size = Pt(10)
        run2.font.name = 'Calibri'

    # Institution on next line
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(6)
    run = p2.add_run(institution)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'

def create_resume():
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # ===== HEADER =====
    # Name - large, centered, blue
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Charles Martin, Ed.D.")
    run.font.size = Pt(22)
    run.font.name = 'Calibri'
    run.font.color.rgb = BLUE_ACCENT
    p.paragraph_format.space_after = Pt(0)

    # Contact info - centered, smaller
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("charlemsartinedd@gmail.com | (323) 632-2071 | San Diego, CA")
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.font.color.rgb = GRAY

    # Horizontal line
    add_horizontal_line(doc)

    # ===== SUMMARY =====
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    summary = (
        "Senior Instructional Designer and eLearning Lead with deep experience building complex "
        "training ecosystems for government, healthcare, and K-12/HE partners. Founder of Alexandria's "
        "Design, partnering with Cell Collective, Eccalon, and UCLA Health to deliver scalable, accessible "
        "learning solutions while integrating AI/automation into design and production workflows."
    )
    run = p.add_run(summary)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'

    # ===== WORK EXPERIENCE =====
    add_section_header(doc, "WORK EXPERIENCE")

    # Job 1: Cell Collective / ModelIt!
    add_job_entry(
        doc,
        "Vice President, Educational Innovation",
        "Cell Collective / ModelIt!",
        "2022 – Present",
        [
            "Lead instructional design and innovation strategy for interactive biological modeling experiences, partnering with university researchers to build NGSS-aligned lessons.",
            "Oversee vision and roadmap for story-based, gamified learning sequences that blend scientific modeling, narrative, and formative assessment.",
            "Guide development of teacher-facing materials, rubrics, and implementation guides supporting classroom adoption at scale.",
            "Collaborate on data tracking and research requirements for SBIR-style studies, ensuring curriculum assets align with research protocols."
        ]
    )

    # Job 2: Eccalon
    add_job_entry(
        doc,
        "Senior Instructional Designer & eLearning Lead",
        "Eccalon",
        "2023 – Present",
        [
            "Lead instructional designer for the 10+ module Mentor Protégé Program (MPP) training ecosystem, including MPP Program Manager Training and Portal courses.",
            "Design, update, and refine multi-module eLearning experiences in Articulate Rise and Storyline, incorporating SME and government feedback.",
            "Contribute to updates for DAU ACQ 0690 and related OSBP/APEX Accelerators trainings, strengthening interactivity and learner engagement.",
            "Ensure all deliverables comply with OSBP and APEX Accelerators standards for accessibility, SCORM packaging, and quality."
        ]
    )

    # Job 3: UCLA Health
    add_job_entry(
        doc,
        "Instructional Designer / eLearning Specialist",
        "UCLA Health",
        "2022 – Present",
        [
            "Design and develop eLearning for clinical and non-clinical audiences, including ACEs Aware and healthcare-focused training initiatives.",
            "Translate complex clinical content into clear learning objectives, scenarios, and microlearning units using Articulate 360.",
            "Partner with physicians, clinicians, and operational leaders to validate accuracy and align content with organizational priorities.",
            "Support continuous improvement by updating modules, standardizing templates, and incorporating stakeholder feedback into iterative redesigns."
        ]
    )

    # ===== SKILLS =====
    add_section_header(doc, "SKILLS")

    # Technical Skills
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("TECHNICAL SKILLS: ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run = p.add_run(
        "Articulate 360 (Rise 360, Storyline 360), SCORM Packaging, LMS Platforms, "
        "Adobe Creative Cloud, n8n, Make, Zapier, Monday.com, Notion, Google Workspace, Microsoft 365"
    )
    run.font.size = Pt(10)
    run.font.name = 'Calibri'

    # Other Skills
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("OTHER SKILLS: ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run = p.add_run(
        "Curriculum Design (ADDIE, SAM, NGSS), SME Facilitation, Stakeholder Management, "
        "AI-Assisted Content Development, Multi-Stakeholder Project Management, Bilingual (English/Spanish)"
    )
    run.font.size = Pt(10)
    run.font.name = 'Calibri'

    # ===== EDUCATION =====
    add_section_header(doc, "EDUCATION")

    add_education_entry(
        doc,
        "Master's in Technology Leadership",
        "Brown University",
        in_progress=True
    )

    add_education_entry(
        doc,
        "Ed.D., Educational Technology",
        "University of Florida"
    )

    add_education_entry(
        doc,
        "M.Ed., Education",
        "Georgia State University"
    )

    add_education_entry(
        doc,
        "B.A., History / Political Science",
        "Georgia State University"
    )

    add_education_entry(
        doc,
        "A.A., Spanish",
        "Foothill College"
    )

    # Save the document
    output_path = os.path.join(os.path.dirname(__file__), "..", "Charles_Martin_Resume.docx")
    output_path = os.path.abspath(output_path)
    doc.save(output_path)
    print(f"Resume saved to: {output_path}")
    return output_path

if __name__ == "__main__":
    create_resume()
